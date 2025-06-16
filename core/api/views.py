from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework import status, generics, viewsets
from rest_framework.permissions import IsAuthenticated, AllowAny, BasePermission
from rest_framework.exceptions import PermissionDenied
from rest_framework_simplejwt.views import (
    TokenObtainPairView,
    TokenRefreshView,
)
from rest_framework.decorators import action

from .serializers import TestGenerationRequestSerializer
from .services.gigachat_service import generate_test_text
from django.utils import timezone
from datetime import datetime

from .serializers import *
from main.models import *

from django.db import transaction, models
from django.http import HttpResponse
from django.shortcuts import get_object_or_404
import json

from .api_docs import *

from docx import Document
from io import BytesIO

@REGISTER_SCHEMA
class RegisterView(APIView):
    def post(self, request):
        serializer = RegisterSerializer(data=request.data)
        if serializer.is_valid():
            serializer.save()
            return Response({'message': 'User registered successfully'}, status=status.HTTP_201_CREATED)
        return Response(serializer.errors, status=status.HTTP_400_BAD_REQUEST)

@AUTH_SCHEMA
class DecoratedTokenObtainPairView(TokenObtainPairView):
    def post(self, request, *args, **kwargs):
        return super().post(request, *args, **kwargs)

@TOKEN_REFRESH_SCHEMA
class DecoratedTokenRefreshView(TokenRefreshView):
    def post(self, request, *args, **kwargs):
        return super().post(request, *args, **kwargs)

@CLASSES_SCHEMA
class ClassroomViewSet(viewsets.ModelViewSet):
    serializer_class = ClassroomSerializer
    permission_classes = [IsAuthenticated]
    http_method_names = ['get', 'post', 'put', 'delete']

    def get_queryset(self):
        return Classroom.objects.filter(owner=self.request.user)

    def create(self, request, *args, **kwargs):
        data = request.data
        if isinstance(data, list):
            # массив классов
            serializer = self.get_serializer(data=data, many=True)
        else:
            serializer = self.get_serializer(data=data)

        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    def perform_create(self, serializer):
        serializer.save(owner=self.request.user)

@TESTS_SCHEMA
class TestViewSet(viewsets.ModelViewSet):
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        return Test.objects.filter(created_by=self.request.user)

    def get_serializer_class(self):
        if self.action == 'list':
            return TestListSerializer
        return TestDetailSerializer

    def perform_create(self, serializer):
        serializer.save(created_by=self.request.user)

    def get_object(self):
        test = super().get_object()
        if test.created_by != self.request.user:
            raise PermissionDenied("Вы не имеете доступа к этому тесту.")
        return test

    @action(detail=True, methods=['post'], url_path='clone')
    def clone(self, request, pk=None):
        original_test = self.get_object()

        # Создаем клон теста
        clone_title = f"{original_test.title} (клон)"
        cloned_test = Test.objects.create(
            title=clone_title,
            created_by=request.user
        )

        # Клонируем все вопросы и ответы
        for question in original_test.questions.all():
            cloned_question = Question.objects.create(
                test=cloned_test,
                text=question.text,
                question_type=question.question_type
            )

            for answer in question.answers.all():
                Answer.objects.create(
                    question=cloned_question,
                    text=answer.text,
                    is_correct=answer.is_correct
                )

        return Response({
            'id': cloned_test.id,
            'title': cloned_test.title,
            'question_count': cloned_test.questions.count()
        }, status=status.HTTP_201_CREATED)

@TEST_LAUNCH_SCHEMA
class TestLaunchViewSet(viewsets.ModelViewSet):
    serializer_class = TestLaunchSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        queryset = TestLaunch.objects.filter(test__created_by=self.request.user)

        # Обновим статусы просроченных сессий
        expired = queryset.filter(expires_at__lte=timezone.now(), is_active=True)
        expired.update(is_active=False)

        return queryset.order_by('-launched_at')

    def perform_create(self, serializer):
        test = serializer.validated_data['test']
        if test.created_by != self.request.user:
            raise PermissionDenied("Вы не являетесь владельцем этого теста.")
        serializer.save()

    def perform_update(self, serializer):
        instance = serializer.instance
        if instance.test.created_by != self.request.user:
            raise PermissionDenied("Вы не являетесь владельцем этого теста.")
        serializer.save()

    def perform_destroy(self, instance):
        if instance.test.created_by != self.request.user:
            raise PermissionDenied("Вы не являетесь владельцем этого теста.")
        instance.delete()

@STUDENTS_SCHEMA
class StudentViewSet(viewsets.ModelViewSet):
    serializer_class = StudentSerializer
    permission_classes = [IsAuthenticated]

    # Получение всех студентов по классу
    def get_queryset(self):
        classroom_id = self.request.query_params.get('classroom_id')
        queryset = Student.objects.filter(classroom__owner=self.request.user)
        if classroom_id:
            queryset = queryset.filter(classroom_id=classroom_id)
        return queryset

    def create(self, request, *args, **kwargs):
        data = request.data
        if isinstance(data, list):
            serializer = self.get_serializer(data=data, many=True, context={'request': request})
        else:
            serializer = self.get_serializer(data=data, context={'request': request})

        serializer.is_valid(raise_exception=True)
        self.perform_create(serializer)
        return Response(serializer.data, status=status.HTTP_201_CREATED)

    def perform_create(self, serializer):
        if isinstance(serializer.validated_data, list):
            for item in serializer.validated_data:
                item['student_id'] = uuid.uuid4().hex[:12]
            serializer.save()
        else:
            serializer.save()

@SUBMIT_ANSWERS_SCHEMA
# Отправка выбранных вариантов ответов на вопрос
class SubmitAnswersStudentView(APIView):
    def post(self, request):
        serializer = SubmitAnswersStudentSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        student_id = serializer.validated_data["student_id"]
        test_launch_id = serializer.validated_data["test_launch_id"]
        answers = serializer.validated_data["answers"]

        student = get_object_or_404(Student, student_id=student_id)
        test_launch = get_object_or_404(TestLaunch, id=test_launch_id)
        test = test_launch.test  # получаем сам тест из сессии

        if StudentTestResult.objects.filter(student=student, test_launch=test_launch).exists():
            return Response({"error": "Этот студент уже завершил тест."}, status=403)

        if not test_launch.classrooms.filter(students=student).exists():
            return Response({"error": "Этот студент не имеет доступа к данному тесту."}, status=403)

        with transaction.atomic():
            answers_by_question = {ans["question"]: ans for ans in answers}
            correct_count = 0
            total_checked = 0

            for question in test.questions.all():
                answer_data = answers_by_question.get(question.id)
                selected_answers = answer_data.get("selected_answers", []) if answer_data else []

                student_answer = StudentAnswer.objects.create(
                    student=student,
                    question=question,
                    test_launch=test_launch,
                    is_checked=False,
                    is_correct=False
                )
                if selected_answers:
                    student_answer.selected_answers.set(selected_answers)

                # Автоматическая проверка:
                if question.question_type in ["one", "multiple", "true_false"]:
                    correct_ids = set(question.answers.filter(is_correct=True).values_list("id", flat=True))
                    selected_ids = set(student_answer.selected_answers.values_list("id", flat=True))
                    is_correct = correct_ids == selected_ids
                    if selected_answers:
                        student_answer.is_checked = True
                    student_answer.is_correct = is_correct
                    student_answer.save()

                    total_checked += 1
                    if is_correct:
                        correct_count += 1
                else:
                    student_answer.save()

            score = (correct_count / total_checked) * 100 if total_checked else 0

            StudentTestResult.objects.create(
                student=student,
                test_launch=test_launch,
                score=score,
                completed_at=timezone.now()
            )

        return Response({"message": "Ответы сохранены", "score": score})

@SESSION_RESULTS_SCHEMA
# Вывод результатов по сессии
class TestLaunchResultView(generics.ListAPIView):
    serializer_class = StudentTestResultSerializer
    permission_classes = [IsAuthenticated]

    def get_queryset(self):
        launch_id = self.kwargs.get('launch_id')
        test_launch = get_object_or_404(TestLaunch, id=launch_id)

        if test_launch.test.created_by != self.request.user:
            return StudentTestResult.objects.none()

        return StudentTestResult.objects.filter(test_launch=test_launch)
@STUDENT_TEST_ANSWERS_SCHEMA
# Получение информации (выбранных ответов) о прохождении теста учеником
class StudentTestAnswersView(APIView):
    def get(self, request, student_id, launch_id):
        student = get_object_or_404(Student, student_id=student_id)
        test_launch = get_object_or_404(TestLaunch, id=launch_id)

        # Проверка доступа
        if not test_launch.classrooms.filter(students=student).exists():
            return Response({"error": "Студент не имеет доступа к данной сессии теста."}, status=403)

        answers = StudentAnswer.objects.filter(
            student=student,
            question__test=test_launch.test
        )

        result = StudentTestResult.objects.filter(
            student=student,
            test_launch=test_launch
        ).first()

        serializer = StudentAnswerDetailSerializer(answers, many=True)

        return Response({
            "score": result.score if result else None,
            "completed_at": result.completed_at if result else None,
            "answers": serializer.data
        })

def check_student_answers(student, test_launch, answers):
    correct_count = 0
    total_checked = 0

    with transaction.atomic():
        for ans in answers:
            question_id = ans["question"]
            selected_answers = ans.get("selected_answers", [])

            question = Question.objects.get(id=question_id)
            student_answer = StudentAnswer.objects.create(
                student=student,
                test_launch=test_launch,
                question=question,
            )
            student_answer.selected_answers.set(selected_answers)

            if question.question_type in ['one', 'multiple', 'true_false']:
                correct_ids = set(question.answers.filter(is_correct=True).values_list('id', flat=True))
                selected_ids = set(student_answer.selected_answers.values_list('id', flat=True))
                is_correct = correct_ids == selected_ids
                student_answer.is_checked = True
                student_answer.is_correct = is_correct

                total_checked += 1
                if is_correct:
                    correct_count += 1

            student_answer.save()

        score = (correct_count / total_checked) * 100 if total_checked else 0

        StudentTestResult.objects.create(
            student=student,
            test_launch=test_launch,
            score=score,
            completed_at=timezone.now()
        )

        return score

# class MarkTextAnswerView(generics.UpdateAPIView):
#     queryset = StudentAnswer.objects.filter(question__question_type='text')
#     serializer_class = TextAnswerCheckSerializer
#     permission_classes = [IsAuthenticated]

@GENERATE_TEST_SCHEMA
class GenerateTestView(APIView):
    # permission_classes = [IsAuthenticated]
    permission_classes = [AllowAny] # Доступно для обоих типов пользователей

    def post(self, request):
        serializer = TestGenerationRequestSerializer(data=request.data)
        serializer.is_valid(raise_exception=True)
        data = serializer.validated_data

        prompt = self.build_prompt(data)
        ai_response = generate_test_text(prompt)

        # Попробуем извлечь JSON внутри ````json ... ```
        try:
            # Очистка от markdown-обёртки ```
            json_text = self.extract_json(ai_response)
            questions_data = json.loads(json_text)
        except Exception as e:
            return Response({"error": f"Ошибка парсинга JSON: {str(e)}",
                                  "AI_answer": ai_response}, status=status.HTTP_400_BAD_REQUEST)

        # Создание теста и вопросов
        # UPD: Формирование валидного JSON для фронта, с возможностью использования реализованного механизма сохранения теста
        try:
            # with transaction.atomic():
            #     test = Test.objects.create(
            #         title=f"«{data['topic']}»",
            #         created_by=request.user
            #     )
            #     self.create_questions(test, questions_data)
            #
            # return Response({
            #     'message': 'Тест успешно создан',
            #     'test_id': test.id,
            #     'answer_AI': ai_response
            # })
            generated = {
                "title": f"«{data['topic']}»",
                "questions": []
            }
            for q in questions_data:
                answers = []
                correct = set(q['correct_answers'])

                for option in q['options']:
                    answers.append({
                        "text": option,
                        "is_correct": option in correct
                    })

                generated["questions"].append({
                    "text": q['question'],
                    "question_type": q['type'],
                    "answers": answers
                })
            return Response(generated, status=200)
        except Exception as e:
            return Response({"error": f"Ошибка обработки данных: {str(e)}"},
                            status=status.HTTP_500_INTERNAL_SERVER_ERROR)

    def extract_json(self, raw: str) -> str:
        """
        Убирает возможные ```json и ``` из ответа.
        """
        if raw.startswith("```json"):
            raw = raw[len("```json"):].strip()
        if raw.endswith("```"):
            raw = raw[:-3].strip()
        return raw

    # def create_questions(self, test, questions_data):
    #     for q in questions_data:
    #         q_type = q['type']
    #         q_text = q['question']
    #         question = Question.objects.create(
    #             test=test,
    #             text=q_text,
    #             question_type=q_type
    #         )
    #
    #         if q_type in ['one', 'multiple', 'true_false']:
    #             options = q['options']
    #             correct = set(q['correct_answers'])
    #
    #             for option in options:
    #                 Answer.objects.create(
    #                     question=question,
    #                     text=option,
    #                     is_correct=option in correct
    #                 )

    def build_prompt(self, data):
        """
        Строит промпт для GigaChat на основе темы и структуры типов вопросов.

        :param topic: тема теста
        :param types: словарь с количеством вопросов каждого типа, например:
                      {"one": 3, "multiple": 2, "true_false": 1}
        :return: текст запроса
        """
        topic = data['topic']
        distribution = data['type_distribution']
        one_count = distribution.get("one", 0)
        multiple_count = distribution.get("multiple", 0)
        true_false_count = distribution.get("true_false", 0)

        prompt = f"""Сгенерируй вопросы по теме «{topic}», СТРОГО придерживаясь следующего формата:
Верни список вопросов в виде JSON-массива, где каждый вопрос — это объект со следующими полями:
- "type": тип вопроса ("one", "multiple", "true_false");
- "question": текст вопроса;
- "options": массив ТОЛЬКО ЛИБО из 4 вариантов ответа (ТОЛЬКО для "one" и "multiple"), ЛИБО из 2 вариантов ответов (ТОЛЬКО для true_false)";
- "correct_answers": массив правильных ответов (для "one" и "true_false" ТОЛЬКО один, а для "multiple" МОГУТ БЫТЬ все правильные или несколько).
Примеры:
[
  {{
    "type": "one",
    "question": "Что такое переменная в программировании?",
    "options": ["Число", "Хранилище данных", "Функция", "Цикл"],
    "correct_answers": ["Хранилище данных"]
  }},
  {{
    "type": "multiple",
    "question": "Что из ниже перечисленного — структуры управления потоком?",
    "options": ["if", "for", "int", "while"],
    "correct_answers": ["if", "for", "int"]
  }},
  {{
    "type": "true_false",
    "question": "Используются ли переменные в структурах управления потоком?",
    "options": ["Правда", "Ложь"],
    "correct_answers": ["Правда"]
  }}
]
Сгенерируй РОВНО:
- {one_count} вопрос(ов) типа "one";
- {multiple_count} вопрос(ов) типа "multiple";
- {true_false_count} вопрос(ов) типа "true_false".
Верни только валидный JSON-массив вопросов. Никаких пояснений или форматирования вне JSON. Ответ должен быть пригоден для автоматического парсинга.
"""
        return prompt


class ExportWordView(APIView):
    permission_classes = []  # доступно всем

    def post(self, request):
        try:
            data = request.data
            mode = data.get("mode", "test")  # 'test' или 'answers'
            title = data.get("title", "Без названия")
            questions = data.get("questions", [])

            doc = Document()
            doc.add_heading(title, level=0)

            for q in questions:
                doc.add_paragraph(f"{q['text']}", style='List Number')

                for idx, a in enumerate(q["answers"], start=1):
                    if mode == "answers":
                        text = f" {idx}) {a['text']} {'(+)' if a['is_correct'] else ''}"
                    else:
                        text = f" {idx}) {a['text']}"
                    doc.add_paragraph(text, style='BodyText')

                doc.add_paragraph("")  # пустая строка между вопросами

            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)

            response = HttpResponse(
                buffer.getvalue(),
                content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document'
            )
            filename = f"{'Тест' if mode == 'test' else 'Ответы'}_{title}.docx"
            response['Content-Disposition'] = f'attachment; filename="{filename}"'
            return response

        except Exception as e:
            return Response({"error": str(e)}, status=status.HTTP_400_BAD_REQUEST)
